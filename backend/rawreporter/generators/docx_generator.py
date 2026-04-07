import logging
from io import BytesIO
from uuid import UUID

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from fastapi import HTTPException
from sqlalchemy.ext.asyncio import AsyncSession

from rawreporter.generators.context_builder import build_report_context, SEVERITY_ORDER
from rawreporter.generators.template_loader import (
    get_template_path,
    inject_cover_page_bookmarks,
    load_template_document,
)
from rawreporter.services.report_service import validate_report_for_generation

logger = logging.getLogger(__name__)

SEVERITY_COLOURS = {
    "critical": RGBColor(0x7C, 0x3A, 0xED),   # purple
    "high": RGBColor(0xDC, 0x26, 0x26),        # red
    "medium": RGBColor(0xD9, 0x77, 0x06),      # amber
    "low": RGBColor(0x25, 0x63, 0xEB),         # blue
    "informational": RGBColor(0x6B, 0x72, 0x80),  # gray
}

SEVERITY_LABELS = {
    "critical": "Critical",
    "high": "High",
    "medium": "Medium",
    "low": "Low",
    "informational": "Informational",
}

REF_LABELS = {
    "cve": "CVE",
    "cwe": "CWE",
    "cisa": "CISA Advisory",
    "nist": "NIST Reference",
    "nvd": "NVD Entry",
    "manufacturer": "Manufacturer Advisory",
}

REF_ORDER = ["cve", "cwe", "cisa", "nist", "nvd", "manufacturer"]


async def generate_docx(report_id: UUID, session: AsyncSession) -> bytes:
    """Main entry point — validates, builds context, generates DOCX bytes."""

    # 1. Validate (raises 422 if blocking findings)
    await validate_report_for_generation(report_id, session)

    # 2. Build context (raises 404 if report not found)
    context = await build_report_context(report_id, session)

    # 3. Get engagement type
    engagement_type = context["engagement"]["type"]

    # 4. Check template exists
    template_path = get_template_path(engagement_type)
    if template_path is None:
        raise HTTPException(
            status_code=422,
            detail={
                "detail": (
                    "No document template uploaded for this report type. "
                    "An Admin must upload a .docx template before reports can be generated."
                ),
                "engagement_type": engagement_type,
                "upload_path": "/settings/document-templates",
            },
        )

    # 5. Load template
    doc = load_template_document(engagement_type)

    # 6. Inject cover page bookmarks
    doc = inject_cover_page_bookmarks(doc, context)

    # 7. Build body content
    build_document_body(doc, context)

    # 8. Save to buffer and return bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_document_body(doc: Document, context: dict) -> None:
    """Appends all report body content to the document after the cover page."""

    # Page break to separate cover from body
    doc.add_page_break()

    # Word TOC field
    _add_toc(doc)
    doc.add_page_break()

    # Render sections
    for section in context["sections"]:
        if section["is_findings_section"]:
            render_findings_section(doc, context)
            continue

        if section["section_type"] == "report_title":
            continue  # title injected via bookmark

        # Section heading (Heading 1)
        doc.add_heading(section["title"], level=1)

        # Body text
        body = section["body_text"]
        if body and body.strip():
            for line in body.split("\n"):
                if line.strip():
                    para = doc.add_paragraph(line)
                    para.style = doc.styles["Normal"]
        else:
            para = doc.add_paragraph("[Section not completed]")
            run = para.runs[0] if para.runs else para.add_run("[Section not completed]")
            if not para.runs:
                para.clear()
                run = para.add_run("[Section not completed]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph("")  # spacing


def _add_toc(doc: Document) -> None:
    """Insert a Word TOC field paragraph."""
    para = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), ' TOC \\o "1-3" \\h \\z \\u ')
    run_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = '[Right-click and select "Update Field" to generate table of contents]'
    run_elem.append(t_elem)
    fld.append(run_elem)
    para._p.append(fld)


def render_findings_section(doc: Document, context: dict) -> None:
    """Render the full findings section with summary table and findings by severity."""

    # Heading
    doc.add_heading("Findings", level=1)

    # Summary table
    summary = context["findings_summary"]
    by_sev = context["findings_by_severity"]

    # Collect rows for severities with findings
    sev_rows = [
        (sev.value, SEVERITY_LABELS[sev.value], summary[sev.value])
        for sev in SEVERITY_ORDER
        if summary.get(sev.value, 0) > 0
    ]

    if sev_rows:
        total = summary.get("total", 0)
        table = doc.add_table(rows=1 + len(sev_rows) + 1, cols=2)
        table.style = "Table Grid"

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(1.5)

        # Header row
        hdr = table.rows[0]
        for cell, text in zip(hdr.cells, ["Severity", "Count"]):
            cell.text = text
            run = cell.paragraphs[0].runs[0]
            run.bold = True

        # Shade header (blue)
        _shade_cell(hdr.cells[0], "4472C4")
        _shade_cell(hdr.cells[1], "4472C4")

        # Data rows
        for i, (sev_val, sev_label, count) in enumerate(sev_rows):
            row = table.rows[i + 1]
            row.cells[0].text = sev_label
            row.cells[1].text = str(count)
            colour = SEVERITY_COLOURS.get(sev_val)
            if colour:
                for cell in row.cells:
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = colour

        # Total row
        total_row = table.rows[-1]
        total_row.cells[0].text = "Total"
        total_row.cells[1].text = str(total)
        for cell in total_row.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

    doc.add_paragraph("")

    # Findings by severity
    for sev in SEVERITY_ORDER:
        sev_val = sev.value
        findings = by_sev.get(sev_val)
        if not findings:
            continue

        count = len(findings)
        label = SEVERITY_LABELS[sev_val]
        doc.add_heading(f"{label} Findings ({count})", level=2)

        for finding in findings:
            render_finding(doc, finding)


def render_finding(doc: Document, finding: dict) -> None:
    """Render one finding with all sub-sections."""
    sev_val = finding["severity"]
    sev_colour = SEVERITY_COLOURS.get(sev_val, RGBColor(0, 0, 0))
    sev_label = SEVERITY_LABELS.get(sev_val, sev_val).upper()

    # Finding title as Heading 2 + severity label
    heading = doc.add_heading(finding["title"], level=2)
    tab_run = heading.add_run(f"\t{sev_label}")
    tab_run.bold = True
    tab_run.font.color.rgb = sev_colour

    # Placement override warning
    if finding.get("is_placement_override"):
        warn_para = doc.add_paragraph()
        warn_run = warn_para.add_run(
            f"\u26A0  Placement Override \u2014 Severity: {SEVERITY_LABELS.get(sev_val, sev_val)}"
        )
        warn_run.italic = True
        warn_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

        justification = finding.get("override_justification") or ""
        just_para = doc.add_paragraph()
        just_run = just_para.add_run(f"Justification: {justification}")
        just_run.italic = True
        just_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    # Summary
    doc.add_heading("Summary", level=3)
    summary = finding.get("summary")
    if summary and summary.strip():
        para = doc.add_paragraph(summary)
        para.style = doc.styles["Normal"]
    else:
        _add_placeholder(doc, "[No summary provided]")

    # Observation
    doc.add_heading("Observation", level=3)
    observation = finding.get("observation")
    if observation and observation.strip():
        para = doc.add_paragraph(observation)
        para.style = doc.styles["Normal"]
    else:
        _add_placeholder(doc, "[No observation recorded]")

    # Recommendation
    doc.add_heading("Recommendation", level=3)
    recommendation = finding.get("recommendation")
    if recommendation and recommendation.strip():
        para = doc.add_paragraph(recommendation)
        para.style = doc.styles["Normal"]
    else:
        _add_placeholder(doc, "[No recommendation provided]")

    # Remediation steps (only if enabled)
    if finding.get("remediation_steps_enabled") and finding.get("remediation_steps"):
        doc.add_heading("Remediation Steps", level=3)
        para = doc.add_paragraph(finding["remediation_steps"])
        para.style = doc.styles["Normal"]

    # CVSS score
    cvss_score = finding.get("cvss_score")
    cvss_vector = finding.get("cvss_vector")
    if cvss_score is not None:
        cvss_para = doc.add_paragraph()
        label_run = cvss_para.add_run("CVSS Score: ")
        label_run.bold = True
        cvss_para.add_run(str(cvss_score))
        if cvss_vector:
            cvss_para.add_run(f" \u2014 {cvss_vector}")

    # Affected systems
    affected = finding.get("affected_systems")
    if affected and affected.strip():
        doc.add_heading("Affected Systems", level=3)
        para = doc.add_paragraph(affected)
        para.style = doc.styles["Normal"]

    # References
    refs = finding.get("references", {})
    has_refs = any(
        refs.get(rt, {}).get("enabled") and refs.get(rt, {}).get("entries")
        for rt in REF_ORDER
    )

    if has_refs:
        doc.add_heading("References", level=3)
        for ref_type in REF_ORDER:
            ref = refs.get(ref_type, {})
            if ref.get("enabled") and ref.get("entries"):
                label = REF_LABELS[ref_type]
                for entry in ref["entries"]:
                    bullet = doc.add_paragraph(style="List Bullet")
                    label_run = bullet.add_run(f"{label}: ")
                    label_run.bold = True
                    bullet.add_run(entry["value"])
                    if entry.get("url"):
                        bullet.add_run(f" \u2014 {entry['url']}")

    # Horizontal rule separator
    _add_horizontal_rule(doc)


def _add_placeholder(doc: Document, text: str) -> None:
    """Add a gray italic placeholder paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _shade_cell(cell, hex_color: str) -> None:
    """Apply a solid background shade to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
    # Make header text white for readability
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_horizontal_rule(doc: Document) -> None:
    """Add a light gray bottom-border paragraph to separate findings."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
