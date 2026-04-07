"""
Phase 4 — Document Generation Tests.

Covers:
  - context_builder: report not found, correct structure, invisible sections,
    disabled remediation steps, disabled references
  - generate_docx: fails without template, fails with blocking finding,
    produces valid DOCX bytes
  - platform_settings API: GET and PUT
  - document_template API: upload rejects non-docx, accepts docx
"""
import os
from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select

from rawreporter.generators.context_builder import build_report_context
from rawreporter.models.client import Client
from rawreporter.models.engagement import Engagement
from rawreporter.models.finding import Finding
from rawreporter.models.finding_reference import FindingReference
from rawreporter.models.library_finding import LibraryFinding
from rawreporter.models.platform_setting import PlatformSetting
from rawreporter.models.report import Report
from rawreporter.models.report_section import ReportSection
from rawreporter.utils.enums import (
    EngagementTypeEnum,
    RefTypeEnum,
    SectionTypeEnum,
    SeverityEnum,
)

TEMPLATE_DIR = "uploads/doc_templates"


# ── Helpers ────────────────────────────────────────────────────────────────

async def _make_client(session, name="Test Corp") -> Client:
    c = Client(name=name)
    session.add(c)
    await session.flush()
    return c


async def _make_engagement(session, client: Client, eng_type=EngagementTypeEnum.vulnerability_assessment) -> Engagement:
    e = Engagement(
        client_id=client.id,
        title="Test Engagement",
        types=[eng_type.value],
    )
    session.add(e)
    await session.flush()
    return e


async def _make_report(session, engagement=None, title="Test Report", types=None) -> Report:
    r = Report(
        title=title,
        types=types or ["vulnerability_assessment"],
        engagement_id=engagement.id if engagement else None,
    )
    session.add(r)
    await session.flush()
    return r


async def _make_section(session, report: Report, section_type=SectionTypeEnum.executive_summary,
                         position=1, is_visible=True, severity_filter=None, body_text=None) -> ReportSection:
    s = ReportSection(
        report_id=report.id,
        section_type=section_type,
        position=position,
        is_visible=is_visible,
        severity_filter=severity_filter,
        title=section_type.value.replace("_", " ").title(),
        body_text=body_text,
    )
    session.add(s)
    await session.flush()
    return s


async def _make_finding(session, report: Report, section: ReportSection,
                         title="Test Finding", severity=SeverityEnum.critical,
                         is_placement_override=False, override_justification=None,
                         remediation_steps_enabled=True, remediation_steps="Fix it.",
                         position=0) -> Finding:
    f = Finding(
        report_id=report.id,
        section_id=section.id,
        title=title,
        summary="Summary text",
        observation="Observation text",
        severity_default=severity,
        severity_effective=severity,
        recommendation="Do something",
        remediation_steps=remediation_steps,
        remediation_steps_enabled=remediation_steps_enabled,
        is_placement_override=is_placement_override,
        override_justification=override_justification,
        position=position,
    )
    session.add(f)
    await session.flush()
    return f


def _make_minimal_docx_bytes() -> bytes:
    """Create the smallest valid .docx and return as bytes."""
    doc = DocxDocument()
    doc.add_paragraph("Cover Page")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _save_test_template(engagement_type: str) -> str:
    """Save a minimal docx template and return the file path."""
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    path = os.path.join(TEMPLATE_DIR, f"{engagement_type}.docx")
    with open(path, "wb") as f:
        f.write(_make_minimal_docx_bytes())
    return path


def _remove_test_template(engagement_type: str) -> None:
    path = os.path.join(TEMPLATE_DIR, f"{engagement_type}.docx")
    if os.path.exists(path):
        os.remove(path)


# ── context_builder tests ──────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_context_builder_report_not_found(session):
    """build_report_context raises 404 for unknown report_id."""
    from fastapi import HTTPException
    with pytest.raises(HTTPException) as exc_info:
        await build_report_context(uuid4(), session)
    assert exc_info.value.status_code == 404


@pytest.mark.asyncio
async def test_context_builder_correct_structure(session):
    """Context has correct findings_summary and section counts for a VA report."""
    c = await _make_client(session)
    eng = await _make_engagement(session, c)
    report = await _make_report(session, eng)

    # Create visible sections including a findings container and severity sub-sections
    title_sec = await _make_section(session, report, SectionTypeEnum.report_title, position=1)
    exec_sec = await _make_section(session, report, SectionTypeEnum.executive_summary, position=2,
                                    body_text="Executive summary text.")
    findings_sec = await _make_section(session, report, SectionTypeEnum.findings, position=6)
    crit_sec = await _make_section(session, report, SectionTypeEnum.critical_findings, position=7,
                                    severity_filter=SeverityEnum.critical)
    high_sec = await _make_section(session, report, SectionTypeEnum.high_findings, position=8,
                                    severity_filter=SeverityEnum.high)
    med_sec = await _make_section(session, report, SectionTypeEnum.medium_findings, position=9,
                                   severity_filter=SeverityEnum.medium)
    low_sec = await _make_section(session, report, SectionTypeEnum.low_findings, position=10,
                                   severity_filter=SeverityEnum.low)
    info_sec = await _make_section(session, report, SectionTypeEnum.informational, position=11,
                                    severity_filter=SeverityEnum.informational)
    close_sec = await _make_section(session, report, SectionTypeEnum.closing, position=12)

    # Add 2 critical findings, 1 high
    f1 = await _make_finding(session, report, crit_sec, title="Crit 1", severity=SeverityEnum.critical, position=0)
    f2 = await _make_finding(session, report, crit_sec, title="Crit 2", severity=SeverityEnum.critical, position=1)
    f3 = await _make_finding(session, report, high_sec, title="High 1", severity=SeverityEnum.high, position=0)

    ctx = await build_report_context(report.id, session)

    assert ctx["findings_summary"]["critical"] == 2
    assert ctx["findings_summary"]["high"] == 1
    assert ctx["findings_summary"]["total"] == 3
    assert ctx["client"]["name"] == c.name

    # Critical findings in order by position
    crit_findings = ctx["findings_by_severity"]["critical"]
    assert len(crit_findings) == 2
    assert crit_findings[0]["title"] == "Crit 1"
    assert crit_findings[1]["title"] == "Crit 2"

    # Sections: severity sub-sections are excluded (rendered by render_findings_section).
    # 4 sections: report_title, executive_summary, findings container, closing
    visible_sections = [s for s in ctx["sections"] if True]
    assert len(visible_sections) == 4


@pytest.mark.asyncio
async def test_context_excludes_invisible_sections(session):
    """Invisible sections are not included in context["sections"]."""
    c = await _make_client(session)
    eng = await _make_engagement(session, c)
    report = await _make_report(session, eng)

    await _make_section(session, report, SectionTypeEnum.executive_summary, position=1, is_visible=True)
    await _make_section(session, report, SectionTypeEnum.appendix, position=2, is_visible=False)

    ctx = await build_report_context(report.id, session)
    section_types = [s["section_type"] for s in ctx["sections"]]
    assert "appendix" not in section_types
    assert "executive_summary" in section_types


@pytest.mark.asyncio
async def test_context_excludes_disabled_remediation_steps(session):
    """Findings with remediation_steps_enabled=False have null remediation_steps in context."""
    c = await _make_client(session)
    eng = await _make_engagement(session, c)
    report = await _make_report(session, eng)

    sec = await _make_section(session, report, SectionTypeEnum.critical_findings, position=1,
                               severity_filter=SeverityEnum.critical)
    await _make_finding(
        session, report, sec,
        severity=SeverityEnum.critical,
        remediation_steps_enabled=False,
        remediation_steps="This should be hidden",
    )

    ctx = await build_report_context(report.id, session)
    findings = ctx["findings_by_severity"]["critical"]
    assert len(findings) == 1
    assert findings[0]["remediation_steps"] is None


@pytest.mark.asyncio
async def test_context_excludes_disabled_references(session):
    """Findings with ref_cve_enabled=False have no CVE entries in context."""
    c = await _make_client(session)
    eng = await _make_engagement(session, c)
    report = await _make_report(session, eng)

    sec = await _make_section(session, report, SectionTypeEnum.critical_findings, position=1,
                               severity_filter=SeverityEnum.critical)
    f = await _make_finding(session, report, sec, severity=SeverityEnum.critical)

    # Add a CVE reference but disable it
    f.ref_cve_enabled = False
    ref = FindingReference(
        finding_id=f.id,
        ref_type=RefTypeEnum.cve,
        identifier="CVE-2024-1234",
        is_visible=True,
    )
    session.add(ref)
    await session.flush()

    ctx = await build_report_context(report.id, session)
    findings = ctx["findings_by_severity"]["critical"]
    assert len(findings) == 1
    cve_ref = findings[0]["references"]["cve"]
    assert cve_ref["enabled"] is False
    assert cve_ref["entries"] == []


# ── generate_docx tests ────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_generate_fails_without_template(session):
    """generate_docx raises 422 when no template has been uploaded."""
    from fastapi import HTTPException
    from rawreporter.generators.docx_generator import generate_docx

    c = await _make_client(session)
    eng = await _make_engagement(session, c)
    report = await _make_report(session, eng)

    # Ensure no template file exists for this type
    _remove_test_template("vulnerability_assessment")

    with pytest.raises(HTTPException) as exc_info:
        await generate_docx(report.id, session)
    assert exc_info.value.status_code == 422
    detail = exc_info.value.detail
    assert "No document template" in str(detail)


@pytest.mark.asyncio
async def test_generate_fails_with_blocking_finding(session):
    """generate_docx raises 422 when a finding has unresolved placement override."""
    from fastapi import HTTPException
    from rawreporter.generators.docx_generator import generate_docx

    c = await _make_client(session)
    eng = await _make_engagement(session, c)
    report = await _make_report(session, eng)

    sec = await _make_section(session, report, SectionTypeEnum.high_findings, position=1,
                               severity_filter=SeverityEnum.high)
    blocking = await _make_finding(
        session, report, sec,
        title="Blocking Finding",
        severity=SeverityEnum.critical,
        is_placement_override=True,
        override_justification=None,  # no justification = blocks generation
    )

    _save_test_template("vulnerability_assessment")
    try:
        with pytest.raises(HTTPException) as exc_info:
            await generate_docx(report.id, session)
        assert exc_info.value.status_code == 422
        assert "blocking_findings" in exc_info.value.detail
    finally:
        _remove_test_template("vulnerability_assessment")


@pytest.mark.asyncio
async def test_generate_produces_valid_docx(session):
    """generate_docx returns valid DOCX bytes for a complete report."""
    from rawreporter.generators.docx_generator import generate_docx

    c = await _make_client(session)
    eng = await _make_engagement(session, c)
    report = await _make_report(session, eng, title="Full Test Report")

    # Minimal section structure
    title_sec = await _make_section(session, report, SectionTypeEnum.report_title, position=1)
    exec_sec = await _make_section(session, report, SectionTypeEnum.executive_summary, position=2,
                                    body_text="Some executive summary.")
    findings_sec = await _make_section(session, report, SectionTypeEnum.findings, position=6)
    crit_sec = await _make_section(session, report, SectionTypeEnum.critical_findings, position=7,
                                    severity_filter=SeverityEnum.critical)

    # One finding with references
    f = await _make_finding(session, report, crit_sec, severity=SeverityEnum.critical)
    f.ref_cve_enabled = True
    ref = FindingReference(
        finding_id=f.id,
        ref_type=RefTypeEnum.cve,
        identifier="CVE-2024-9999",
        url="https://example.com",
        is_visible=True,
    )
    session.add(ref)
    await session.flush()

    _save_test_template("vulnerability_assessment")
    try:
        result = await generate_docx(report.id, session)
        assert isinstance(result, bytes)
        assert len(result) > 0

        # Must be a valid DOCX (openable by python-docx)
        doc = DocxDocument(BytesIO(result))
        assert len(doc.paragraphs) > 0
    finally:
        _remove_test_template("vulnerability_assessment")


# ── Platform settings API tests ────────────────────────────────────────────

@pytest.mark.asyncio
async def test_platform_setting_firm_name(client, session):
    """PUT platform-settings/firm_name updates the value, GET returns it."""
    # Seed the firm_name setting
    existing = await session.execute(
        select(PlatformSetting).where(PlatformSetting.key == "firm_name")
    )
    if existing.scalar_one_or_none() is None:
        session.add(PlatformSetting(key="firm_name", value=None))
        await session.flush()

    resp = await client.put(
        "/api/v1/platform-settings/firm_name",
        json={"value": "Test Security Firm"},
    )
    assert resp.status_code == 200

    resp = await client.get("/api/v1/platform-settings/")
    assert resp.status_code == 200
    data = resp.json()
    assert data["firm_name"] == "Test Security Firm"


# ── Document template upload API tests ────────────────────────────────────

@pytest.mark.asyncio
async def test_document_template_upload_rejects_non_docx(client):
    """Uploading a non-.docx file returns 422."""
    import io
    txt_content = b"not a docx file"
    resp = await client.post(
        "/api/v1/document-templates/vulnerability_assessment",
        files={"file": ("test.txt", io.BytesIO(txt_content), "text/plain")},
    )
    assert resp.status_code == 422


@pytest.mark.asyncio
async def test_document_template_upload_accepts_docx(client):
    """Uploading a valid .docx file returns 200 and saves the file."""
    import io
    docx_bytes = _make_minimal_docx_bytes()

    resp = await client.post(
        "/api/v1/document-templates/vulnerability_assessment",
        files={"file": ("template.docx", io.BytesIO(docx_bytes),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200
    data = resp.json()
    assert data["engagement_type"] == "vulnerability_assessment"
    assert data["template"] is not None
    assert data["template"]["original_filename"] == "template.docx"

    # Clean up
    _remove_test_template("vulnerability_assessment")
